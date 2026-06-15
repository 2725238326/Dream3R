from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


ROOT = Path(r"E:\Dream3R")
TEMPLATE = ROOT / "结题报告--纪博闻 - 副本 (2).docx"
OUTPUT = ROOT / "reports" / "final" / "Dream3R_结题报告_模板版_20260615.docx"


TITLE = "Dream3R：面向前馈式三维重建的状态条件化候选几何融合方法研究"


TOC_LINES = [
    "摘要 1",
    "一、绪论 1",
    "1.1 研究背景 1",
    "1.2 研究现状 2",
    "1.3 研究内容与目标 2",
    "二、相关理论与基础 3",
    "2.1 前馈式三维重建与点图表示 3",
    "2.2 候选几何融合 3",
    "2.3 状态信息与记忆状态 3",
    "2.4 稀疏注意力、Mamba 与 Slot Attention 4",
    "三、架构设计 4",
    "3.1 总体流程 4",
    "3.2 候选几何分支 4",
    "3.3 状态记忆分支 5",
    "3.4 候选质量分支 5",
    "3.5 智能场景选择 5",
    "3.6 输出结果 5",
    "四、结果分析 6",
    "4.1 实验设置 6",
    "4.2 主结果 6",
    "4.3 状态对照 6",
    "4.4 与单独教师模型的对比 7",
    "4.5 点图样例与完成边界 7",
    "结论 8",
    "参考文献 8",
]


def set_text(paragraph, text: str) -> None:
    paragraph.text = text


def remove_body_from(doc: Document, start_index: int) -> None:
    body = doc._body._element
    children = list(body)
    for child in children[start_index:-1]:
        body.remove(child)


def remove_existing_report_body(doc: Document) -> None:
    """Keep cover/TOC sectioning, then remove the old report body."""

    body = doc._body._element
    children = list(body)
    last_section_break = None
    for idx, child in enumerate(children[:-1]):
        ppr = child.find(qn("w:pPr"))
        if ppr is not None and ppr.find(qn("w:sectPr")) is not None:
            last_section_break = idx
    if last_section_break is None:
        raise RuntimeError("could not find body section break before report body")
    remove_body_from(doc, last_section_break + 1)


def remove_empty_toc_tail(doc: Document, toc_end_para_index: int) -> None:
    """Remove unused static-TOC placeholder paragraphs before the body section."""

    paragraphs = doc.paragraphs
    section_para_index = None
    for idx in range(toc_end_para_index + 1, len(paragraphs)):
        ppr = paragraphs[idx]._element.find(qn("w:pPr"))
        if ppr is not None and ppr.find(qn("w:sectPr")) is not None:
            section_para_index = idx
            break
    if section_para_index is None:
        return

    for idx in range(section_para_index - 1, toc_end_para_index, -1):
        paragraph = paragraphs[idx]
        if paragraph.text.strip():
            continue
        paragraph._element.getparent().remove(paragraph._element)


def add_p(doc: Document, text: str, style: str = "论文正文", align=None):
    p = doc.add_paragraph(text, style=style)
    if align is not None:
        p.alignment = align
    return p


def add_table(doc: Document, rows: list[list[str]], caption: str, table_style) -> None:
    add_p(doc, caption, "表题")
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    if table_style is not None:
        table.style = table_style
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.text = value
            for p in cell.paragraphs:
                p.style = "论文正文"
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r == 0 or c > 0 else WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph("")


def build_body(doc: Document, table_style) -> None:
    add_p(doc, "摘要", "Heading1")
    for text in [
        "本课题围绕前馈式三维重建中的多模型候选融合问题展开。DUSt3R、MASt3R、Fast3R、Spann3R、VGGT 等模型已经可以从图像中直接预测点图或相关三维表示，省去了传统重建中一部分特征匹配和多步优化流程。但这些模型的优势并不一致：有的速度快，有的匹配稳定，有的更适合长序列或多视图场景。只固定使用一个模型，容易在场景变化时出现结果波动。",
        "Dream3R 的做法是把多个前馈式三维重建教师模型的输出先整理成候选几何，再结合状态信息、记忆上下文、候选置信度和冲突分数进行融合。当前结题版本为 Dream3R v1.1.0。该版本在 KITTI 类输入上采用稳定的 v1.0-rc1 分支，在 ETH3D 类输入上采用引入 VGGT-Omega 后的融合分支，最终输出融合点图、置信度和候选权重。",
        "在现有验证记录中，Dream3R v1.1.0 在 KITTI 和 ETH3D 上的 AbsRel 分别为 0.1448 和 0.0570。状态对照中，正常状态在两个数据域上均低于无状态和乱序状态，说明状态信息在当前候选融合流程中起到了作用。本文只围绕已经完成并验证的 Dream3R v1.1.0 展开，不把未通过验证的探索内容写成正式成果。",
        "关键词：前馈式三维重建；Dream3R；候选几何融合；状态条件化；点图",
    ]:
        add_p(doc, text)

    add_p(doc, "一、绪论", "Heading1")
    add_p(doc, "研究背景", "Heading2")
    for text in [
        "三维重建是计算机视觉中的基础问题，目标是从图像、视频或多视图观测中恢复场景的空间结构。传统路线通常依赖特征提取、匹配、相机位姿估计、多视图几何和后处理优化。这条路线解释性强，也有 COLMAP 等成熟工具，但流程较长，对纹理、视角重叠、相机运动和中间参数比较敏感。",
        "前馈式三维重建模型缩短了这一流程。以 DUSt3R、MASt3R、Fast3R、Spann3R、VGGT 等模型为代表的方法，可以在一次或少量前向推理中输出点图、深度、置信度或相关几何信息。它们更像是直接给出几何候选，而不是把重建拆成一长串手工步骤。",
        "问题也很明显。不同 3R 模型在不同输入上的表现不一致：长序列可能出现漂移，动态物体会干扰静态几何假设，部分模型缺少结果自检查，输入帧数增加后推理成本也会上升。课题后期的工作重点因此从“选一个最好模型”转向“把多个模型的候选结果组织起来，再做可验证的融合”。",
    ]:
        add_p(doc, text)

    add_p(doc, "研究现状", "Heading2")
    for text in [
        "传统三维重建以 SfM 和 MVS 为代表，通过图像匹配和几何约束恢复相机与场景结构。这类方法在静态、高重叠、纹理充足的场景中可靠，但在低纹理、动态场景和快速批量实验中不够方便。",
        "学习式三维重建方法把深度估计、多视图匹配或场景表示交给神经网络完成。MVSNet、NeuralRecon、NeRF 等工作从不同角度推进了这个方向，但其中不少方法仍依赖相机参数、深度监督或较长优化时间。",
        "近两年更贴近本课题的是前馈式 3R 模型族。DUSt3R 直接预测点图；MASt3R 强化匹配与三维几何的联系；Fast3R 关注更高效的多图重建；Spann3R 引入空间记忆；VGGT 提供了较强的几何预测能力；CUT3R、Test3R 等工作则从流式推理和几何检查等方向补充这一体系。这些模型给 Dream3R 提供了候选几何来源，也暴露了多模型融合的需求。",
    ]:
        add_p(doc, text)

    add_p(doc, "研究内容与目标", "Heading2")
    for text in [
        "本课题的目标是完成一个边界清楚、能够运行和验证的候选几何融合模型。具体目标包括：整理多个前馈式 3R 教师模型的输出形式；建立统一候选几何输入；设计状态条件化融合模块；在 KITTI 和 ETH3D 上验证主结果和状态对照；把没有通过验证的内容留在实验记录中，不放进正式结论。",
        "结题阶段形成的正式成果是 Dream3R v1.1.0。它接收教师模型给出的候选点图和置信度，结合梦境状态、记忆状态、候选冲突分数和智能场景选择，输出最终点图、置信度和候选权重。",
    ]:
        add_p(doc, text)

    add_p(doc, "二、相关理论与基础", "Heading1")
    add_p(doc, "前馈式三维重建与点图表示", "Heading2")
    for text in [
        "前馈式三维重建把输入图像送入神经网络后，直接得到三维几何相关输出。点图是其中常见的一种表示方式，可以把每个图像位置对应到三维点坐标。相比只输出深度图，点图更容易表达图像像素与空间位置之间的关系，也更适合不同模型结果之间的候选对齐和融合。",
        "Dream3R 当前使用的教师候选主要以点图和置信度形式进入融合模块。置信度不是最终结论，只是候选质量的一个输入信号。最终权重仍由融合模块结合状态和冲突情况计算。",
    ]:
        add_p(doc, text)

    add_p(doc, "候选几何融合", "Heading2")
    for text in [
        "候选几何融合解决的是多个模型结果如何合并的问题。一个样本可能同时得到 Fast3R、MASt3R、Spann3R、VGGT-Omega 等候选点图。不同候选可能在局部细节、整体尺度或稳定性上各有优劣，也可能互相冲突。",
        "如果简单选择单个教师模型，容易丢掉其他模型在局部区域的有效信息；如果直接平均，又可能把明显错误的候选也混进最终结果。Dream3R 采用候选权重的方式，让模型根据置信度、冲突分数和状态信息调整不同候选的贡献。",
    ]:
        add_p(doc, text)

    add_p(doc, "状态信息与记忆状态", "Heading2")
    for text in [
        "梦境状态表示当前重建任务中的状态信号，可以理解为融合时需要参考的当前条件。记忆状态更偏向历史上下文，记录前面窗口、前面视角或已有候选结果带来的信息。二者都不是最终几何，也不替代教师模型。",
        "在 Dream3R 中，状态信息进入候选权重计算过程。为了证明它不是形式上的输入，实验设置了正常状态、无状态和乱序状态三组对照。只有正常状态稳定优于后两者，才说明状态信息在当前融合中有用。",
    ]:
        add_p(doc, text)

    add_p(doc, "稀疏注意力、Mamba 与 Slot Attention", "Heading2")
    for text in [
        "开题阶段关注过几类机制：稀疏注意力用于筛选有效上下文，Mamba 用于建模序列状态变化，Slot Attention 用于整理候选或局部片段。它们在本课题中主要服务于架构设计和模块划分，不单独写成已经完成的核心结果。",
        "最终报告中只保留它们和 Dream3R 的关系：稀疏注意力对应长序列中哪些上下文值得保留；Mamba 对应状态随序列变化的建模思路；Slot Attention 对应候选或局部几何片段的整理思路。正式结果仍以候选几何融合和状态对照为准。",
    ]:
        add_p(doc, text)

    add_p(doc, "三、架构设计", "Heading1")
    add_p(doc, "总体流程", "Heading2")
    for text in [
        "Dream3R v1.1.0 的流程可以概括为：输入图像先由前馈式 3R 教师模型生成候选点图和置信度；候选结果进入候选几何库；状态记忆分支提供当前状态和历史上下文；候选质量分支计算置信度和冲突信息；智能场景选择确定采用已验证的分支；融合模块输出最终点图、置信度和候选权重。",
        "这一路线保留了教师模型的几何能力，也把最后一步从硬选择改成可检查的融合。Dream3R 不负责替代所有教师模型，而是负责在已有候选之间做更稳的组合。",
    ]:
        add_p(doc, text)

    add_p(doc, "候选几何分支", "Heading2")
    for text in [
        "候选几何分支接收不同教师模型的点图输出。当前结题材料中涉及的主要候选包括 Fast3R、MASt3R、Spann3R 和 VGGT-Omega。不同数据域可用的候选数量不完全一样，KITTI 分支主要使用 Fast3R、MASt3R、Spann3R，ETH3D 分支加入 VGGT-Omega。",
        "候选几何分支的任务是把不同来源的点图和置信度整理成统一张量，供后面的融合模块读取。它不直接决定最终输出。",
    ]:
        add_p(doc, text)

    add_p(doc, "状态记忆分支", "Heading2")
    for text in [
        "状态记忆分支提供梦境状态和记忆状态。梦境状态反映当前窗口的情况，记忆状态反映前序上下文。该分支把远程状态、空间锚点和当前上下文整理后交给融合模块。",
        "这一分支需要靠状态对照来验证。正常状态、无状态和乱序状态的差异，是判断该分支是否有用的主要依据。",
    ]:
        add_p(doc, text)

    add_p(doc, "候选质量分支", "Heading2")
    for text in [
        "候选质量分支处理两个信息：候选置信度和候选冲突。置信度来自教师模型自身输出，冲突分数反映不同候选之间的分歧。二者一起帮助融合模块判断当前候选是否可靠。",
        "在实际结果中，候选之间并不总是一致。某个教师模型可能整体偏好某类场景，也可能在局部出现明显错误。把冲突分数显式送入融合过程，可以减少盲目相信单个候选的情况。",
    ]:
        add_p(doc, text)

    add_p(doc, "智能场景选择", "Heading2")
    for text in [
        "智能场景选择根据已知数据域选择经过验证的分支。KITTI 类输入使用稳定的 v1.0-rc1 分支；ETH3D 类输入使用引入 VGGT-Omega 后的融合分支。",
        "这样处理的原因很直接：不同数据域中更稳定的候选组合不同。把分支选择写清楚，比把所有场景强行塞进同一条融合路线更可靠。",
    ]:
        add_p(doc, text)

    add_p(doc, "输出结果", "Heading2")
    for text in [
        "Dream3R 的输出包括最终点图、最终置信度和候选权重。点图用于表达三维几何结果，置信度用于辅助判断输出可靠性，候选权重用于查看融合过程更依赖哪些教师模型。",
        "这种输出形式便于和单教师模型比较，也便于做点图和点云可视化。结题阶段已经生成了若干真实缓存样本的点图对比图，用于说明 Dream3R 在具体样本上的改进情况。",
    ]:
        add_p(doc, text)

    add_p(doc, "四、结果分析", "Heading1")
    add_p(doc, "实验设置", "Heading2")
    for text in [
        "结果分析围绕 KITTI 和 ETH3D 两个数据域展开。KITTI 更接近室外车载场景，ETH3D 更接近多视图场景。主要指标采用 AbsRel，数值越低表示误差越小。",
        "对照分为三类：一是 Dream3R v1.1.0 与稳定回退版本 v1.0-rc1 的比较；二是正常状态、无状态和乱序状态的比较；三是 Dream3R 与单独教师模型的比较。点图样例基于真实候选缓存样本，用于展示具体输出形态。",
    ]:
        add_p(doc, text)

    add_p(doc, "主结果", "Heading2")
    add_table(
        doc,
        [
            ["模型版本", "KITTI AbsRel", "ETH3D AbsRel", "说明"],
            ["v1.0-rc1", "0.1448", "0.1475", "稳定回退版本"],
            ["Dream3R v1.1.0", "0.1448", "0.0570", "当前正式版本"],
        ],
        "表 4.1 Dream3R 正式版本与稳定回退版本对比",
        table_style,
    )
    for text in [
        "从主结果看，v1.1.0 在 KITTI 上保持了 v1.0-rc1 的结果，在 ETH3D 上误差降到 0.0570。这个结果说明，当前分域融合策略对 ETH3D 更有效，同时没有牺牲 KITTI 上已经稳定的结果。",
    ]:
        add_p(doc, text)

    add_p(doc, "状态对照", "Heading2")
    add_table(
        doc,
        [
            ["数据域", "正常状态", "无状态", "乱序状态"],
            ["KITTI", "0.1448", "0.1553", "0.1521"],
            ["ETH3D", "0.0570", "0.0583", "0.0598"],
        ],
        "表 4.2 Dream3R v1.1.0 状态对照结果",
        table_style,
    )
    add_p(doc, "两个数据域中，正常状态均取得最低误差。KITTI 上正常状态比无状态和乱序状态更低；ETH3D 上正常状态也优于两组对照。这说明状态信息不是只作为输入摆在结构里，而是在当前融合设置下影响了候选权重和最终结果。")

    add_p(doc, "与单独教师模型的对比", "Heading2")
    add_table(
        doc,
        [
            ["数据域", "Fast3R", "MASt3R", "Spann3R", "VGGT-Omega", "Dream3R"],
            ["KITTI", "0.2326", "0.1523", "0.1618", "-", "0.1448"],
            ["ETH3D", "0.1390", "0.1342", "0.1148", "0.0913", "0.0570"],
        ],
        "表 4.3 Dream3R 与单独教师模型 AbsRel 对比",
        table_style,
    )
    for text in [
        "在 KITTI 上，Dream3R 低于该组单教师中的最好结果；在 ETH3D 上，Dream3R 低于 VGGT-Omega 和其他教师模型。这里的结论只针对当前数据和评估设置，不扩展为所有场景下的通用最优。",
        "点图样例也能看到类似趋势。ETH3D 样本 30 中，最好单教师 Spann3R 的 AbsRel 为 0.2534，Dream3R 为 0.1400；ETH3D 样本 35 中，最好单教师 VGGT-Omega 为 0.2178，Dream3R 为 0.1253；KITTI 样本 164 中，最好单教师 MASt3R 为 0.1161，Dream3R 为 0.1064。这些样例用于说明融合结果在具体输入上的改善，不当作新的完整公开评测。",
    ]:
        add_p(doc, text)

    add_p(doc, "点图样例与完成边界", "Heading2")
    add_table(
        doc,
        [
            ["样本", "最好单教师", "单教师 AbsRel", "Dream3R AbsRel", "改善"],
            ["ETH3D 30", "Spann3R", "0.2534", "0.1400", "44.8%"],
            ["ETH3D 35", "VGGT-Omega", "0.2178", "0.1253", "42.5%"],
            ["KITTI 164", "MASt3R", "0.1161", "0.1064", "8.3%"],
        ],
        "表 4.4 真实候选缓存样本的点图对比结果",
        table_style,
    )
    for text in [
        "结题阶段完成了模型流程、候选缓存运行、指标表和点图对比图。没有通过验证的扩展尝试不写入正式结果，只作为后续研究记录。",
        "因此，本报告只把 Dream3R v1.1.0 作为正式成果。后续工作可以继续扩大数据和消融范围，但当前结论不超出已有验证结果。",
    ]:
        add_p(doc, text)

    add_p(doc, "结论", "Heading1")
    for text in [
        "本课题完成了 Dream3R v1.1.0。它是一个面向前馈式三维重建的状态条件化候选几何融合模型，接收多个教师模型的候选点图和置信度，结合状态记忆、候选质量和智能场景选择输出最终点图。",
        "从结果看，Dream3R v1.1.0 在 KITTI 上保持稳定回退版本的 0.1448，在 ETH3D 上达到 0.0570。状态对照中，正常状态优于无状态和乱序状态，说明状态信息在当前融合流程中有实际作用。与单独教师模型相比，Dream3R 在已有对比表和点图样例中也取得了更低误差。",
        "当前工作的边界也很明确。Dream3R v1.1.0 是候选几何融合模型，正式结论只来自已经通过验证的分支和结果。后续如果继续推进，可以扩大评测样本，补充更多消融，并继续改进候选生成和融合策略。",
    ]:
        add_p(doc, text)

    add_p(doc, "参考文献", "Heading1")
    refs = [
        "Schonberger J. L., Frahm J.-M. Structure-From-Motion Revisited[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition. 2016: 4104-4113.",
        "Schonberger J. L., Zheng E., Pollefeys M., Frahm J.-M. Pixelwise View Selection for Unstructured Multi-View Stereo[C]//European Conference on Computer Vision. 2016.",
        "Yao Y., Luo Z., Li S., Fang T., Quan L. MVSNet: Depth Inference for Unstructured Multi-view Stereo[C]//European Conference on Computer Vision. 2018.",
        "Wang S., Leroy V., Cabon Y., Chidlovskii B., Revaud J. DUSt3R: Geometric 3D Vision Made Easy[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. 2024.",
        "Leroy V., Cabon Y., Revaud J. Grounding Image Matching in 3D with MASt3R[C]//European Conference on Computer Vision. 2024.",
        "Yang J., Sax A., Liang K. J., Henaff M., Tang H., Cao A., Chai J., Meier F., Feiszli M. Fast3R: Towards 3D Reconstruction of 1000+ Images in One Forward Pass[EB/OL]. 2024.",
        "Wang H., Agapito L. 3D Reconstruction with Spatial Memory[C]//International Conference on 3D Vision. 2025.",
        "Vaswani A., Shazeer N., Parmar N., Uszkoreit J., Jones L., Gomez A. N., Kaiser L., Polosukhin I. Attention Is All You Need[C]//Advances in Neural Information Processing Systems. 2017.",
        "Gu A., Dao T. Mamba: Linear-Time Sequence Modeling with Selective State Spaces[EB/OL]. 2023.",
        "Locatello F., Weissenborn D., Unterthiner T., et al. Object-Centric Learning with Slot Attention[C]//Advances in Neural Information Processing Systems. 2020.",
        "Dream3R v1.1.0 模型说明、验证记录和结题阶段点图对比材料。",
    ]
    for ref in refs:
        add_p(doc, ref, "参考文献")


def main() -> None:
    doc = Document(TEMPLATE)

    # Keep the cover layout and table intact; only replace the project title.
    set_text(doc.paragraphs[7], TITLE)

    # Replace the static TOC text using the existing TOC paragraph styles.
    for offset, line in enumerate(TOC_LINES, start=16):
        if offset < len(doc.paragraphs):
            set_text(doc.paragraphs[offset], line)
    for offset in range(16 + len(TOC_LINES), 48):
        if offset < len(doc.paragraphs):
            set_text(doc.paragraphs[offset], "")
    remove_empty_toc_tail(doc, 16 + len(TOC_LINES) - 1)

    table_style = deepcopy(doc.tables[1].style) if len(doc.tables) > 1 else None

    # Preserve cover + TOC. Replace all old body content.
    remove_existing_report_body(doc)
    build_body(doc, table_style)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()

