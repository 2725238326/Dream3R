from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.shared import Pt


ROOT = Path(r"E:\Dream3R")
SRC = ROOT / "Dream3R_结题报告.docx"
OUT = ROOT / "Dream3R_结题报告_模块详述版.docx"


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = deepcopy(paragraph._element)
    for child in list(new_p):
        new_p.remove(child)
    paragraph._element.addnext(new_p)
    new_para = paragraph._parent.paragraphs[
        list(paragraph._parent._element.body).index(new_p)
    ] if False else None
    # Re-wrap by walking the document. python-docx has no public paragraph
    # constructor for an arbitrary XML element, so use the private class.
    from docx.text.paragraph import Paragraph

    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def find_para(doc, startswith):
    for p in doc.paragraphs:
        if p.text.strip().startswith(startswith):
            return p
    raise ValueError(f"paragraph not found: {startswith}")


def find_index(doc, startswith):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(startswith):
            return i
    raise ValueError(f"paragraph not found: {startswith}")


def replace_section_until(doc, heading_prefix, next_heading_prefix, new_heading, new_paragraphs):
    start = find_index(doc, heading_prefix)
    end = find_index(doc, next_heading_prefix)
    heading_style = doc.paragraphs[start].style
    if new_heading.startswith("2.6 "):
        heading_style = find_para(doc, "2.5 Mamba").style
    body_style = doc.paragraphs[start + 1].style
    anchor = doc.paragraphs[start]
    anchor.text = new_heading
    anchor.style = heading_style
    for p in list(doc.paragraphs[start + 1:end]):
        delete_paragraph(p)
    current = anchor
    for text in new_paragraphs:
        current = insert_paragraph_after(current, text, body_style)


def add_module_sections(doc):
    body_style = find_para(doc, "Dream3R 的流程可以概括为").style
    h2_style = find_para(doc, "3.2 候选融合模块").style

    # Strengthen 3.1 overview without replacing the whole section.
    p = find_para(doc, "这一路线保留了教师模型的几何能力")
    insert_paragraph_after(
        p,
        "从设计思路看，Dream3R 先保留不同教师模型的长处，再把重点放在候选结果的整理、判断和融合上。这样做的好处是边界清楚：教师模型负责给出可用几何候选，Dream3R 负责判断候选在当前场景下的可信程度，并把多个候选合成一个更稳定的输出。",
        body_style,
    )

    output_heading_index = find_index(doc, "3.3 输出结果")
    anchor = doc.paragraphs[output_heading_index - 1]

    sections = [
        (
            "3.2.4 智能场景选择分支",
            [
                "智能场景选择分支用于判断当前输入更接近哪一类重建条件，并据此选择更合适的融合路径。这里的“场景”主要由候选结果的稳定性、置信度分布、几何冲突程度和输入窗口特点共同决定，后续融合会据此采用更保守的约束，或更充分地利用扩展候选。",
                "这样设计的原因很直接：KITTI 类场景更强调尺度和深度范围的稳定，错误候选一旦被放大，结果会明显变差；ETH3D 类场景结构更复杂，单个教师模型容易在局部区域失误，扩展候选和状态融合的收益更明显。因此，Dream3R v1.1.0 在 KITTI 上沿用稳定的 bounded StatePrior 与残差修正路线，在 ETH3D 上采用 VGGT-Omega 扩展后的 SCF 路线。",
                "在实现上，这一分支给融合模块提供路径选择和约束信息。它决定采用哪一套候选集合、权重约束和融合策略，随后由候选融合模块完成最终点图计算。这样可以说明同一套 Dream3R 框架为什么在不同数据域上使用不同的融合分支。",
            ],
        ),
        (
            "3.2.5 状态条件融合模块",
            [
                "状态条件融合模块是 Dream3R 的核心。它接收候选点图、候选置信度、状态信息和冲突分数，计算每个候选在当前样本中的权重。权重不按固定表格给出，也不只按置信度排序，而是结合候选之间是否一致、状态信息是否支持当前候选、以及不同候选在局部区域的可靠程度来确定。",
                "设计这个模块的动机来自前期对教师模型的观察：单个教师模型并不存在稳定的全局优势。某个模型可能在一类样本上很好，在另一类样本上出现尺度漂移或局部噪声。直接平均会把错误也平均进去，单独选择一个模型又会丢掉其他候选的有效区域。状态条件融合的目标就是在二者之间取一个可控的方案，让不同候选按样本和区域参与融合。",
                "实现时，候选点图会先被整理到统一表示中，随后融合模块计算专家权重，形成基础融合点图。对于稳定分支，Dream3R 使用 StatePrior 给出候选偏好，再用受限残差做小幅修正；对于 ETH3D 分支，系统引入 VGGT-Omega 候选，并通过 SCF 分支完成状态条件融合。这里的残差修正受到候选分歧范围限制，避免模型在候选之外随意生成几何。",
            ],
        ),
        (
            "3.2.6 序列与局部结构辅助模块",
            [
                "稀疏注意力、Mamba 和 Slot Attention 在 Dream3R 中统一作为候选融合的辅助机制使用。稀疏注意力用于筛选有效上下文，避免长序列中无关帧或弱相关区域干扰融合；Mamba 用于建模序列状态变化，帮助状态信息在连续窗口中保持稳定；Slot Attention 用于整理候选或局部片段，把复杂候选拆成更容易比较的局部结构。",
                "这三个模块的共同作用是把输入上下文和候选局部结构整理得更清楚，供融合模块使用。教师候选仍由各 3R 模型生成，最终点图仍由融合模块输出；序列建模和局部整理主要让融合判断更有依据。",
            ],
        ),
        (
            "3.2.7 输出与可解释信息",
            [
                "融合结束后，Dream3R 输出最终点图、最终置信度和候选权重。最终点图用于三维结果展示和误差计算；置信度用于说明哪些区域更可靠；候选权重用于观察模型在当前样本中主要参考了哪些教师结果。",
                "保留候选权重的意义在于，结果不只是一个黑盒点图。答辩或分析时可以结合权重、点图对比和 AbsRel 指标说明 Dream3R 是如何从多个教师候选中得到更稳定结果的。若某个样本效果不好，也能回到候选冲突和权重分配上分析原因。",
            ],
        ),
    ]

    current = anchor
    for heading, paragraphs in sections:
        current = insert_paragraph_after(current, heading, h2_style)
        for text in paragraphs:
            current = insert_paragraph_after(current, text, body_style)


def polish_theory_sections(doc):
    replace_section_until(
        doc,
        "2.4 稀疏注意力",
        "2.5 Mamba",
        "2.4 稀疏注意力",
        [
            "稀疏注意力用于处理长序列或多候选输入中的有效上下文选择。标准注意力会让所有位置两两交互，输入窗口变长后计算量和干扰项都会增加。稀疏注意力只保留部分关键关系，使模型把注意力放在更可能影响当前融合判断的上下文上。",
            "在 Dream3R 中，它的作用可以概括为筛选有效上下文。它不直接输出最终点图，而是帮助状态分支和候选融合模块减少无关信息影响，尤其适合处理多视角窗口中相邻帧、局部区域和候选片段之间的关系。",
        ],
    )
    replace_section_until(
        doc,
        "2.5 Mamba",
        "2.6 Slot Attention",
        "2.5 Mamba",
        [
            "Mamba 属于选择性状态空间模型，适合处理序列状态随时间变化的问题。它通过维护内部状态来传递历史信息，相比完全依赖注意力的做法，在长序列处理中更容易控制计算量。",
            "在 Dream3R 中，Mamba 的位置是序列状态建模。它关注输入窗口前后状态如何变化，而不是直接替代教师模型做几何预测。这样可以把历史上下文压缩成对融合有用的状态信号，供后续候选权重计算使用。",
        ],
    )
    replace_section_until(
        doc,
        "2.6 Slot Attention",
        "三、架构设计",
        "2.6 Slot Attention",
        [
            "Slot Attention 用于把复杂输入整理成若干相对独立的局部片段或候选单元。对 Dream3R 来说，多个教师模型给出的结果并不总是整体一致，局部区域的可靠性也可能不同，因此有必要把候选拆成更便于比较和融合的结构。",
            "本课题中，Slot Attention 的作用是整理候选或局部片段。它帮助模型区分不同候选在局部区域的贡献，为后续权重分配提供更清晰的输入。它在当前工作中作为候选融合的辅助机制使用，不扩展成未验证的独立贡献。",
        ],
    )


def remove_stiff_contrast_phrasing(doc):
    replacements = {
        "前馈式三维重建模型缩短了这一流程。以 DUSt3R、MASt3R、Fast3R、Spann3R、VGGT 等模型为代表的方法，可以在一次或少量前向推理中输出点图、深度、置信度或相关几何信息。它们更像是直接给出几何候选，而不是把重建拆成一长串手工步骤。":
            "前馈式三维重建模型缩短了这一流程。以 DUSt3R、MASt3R、Fast3R、Spann3R、VGGT 等模型为代表的方法，可以在一次或少量前向推理中输出点图、深度、置信度或相关几何信息。它们直接给出几何候选，减少了传统流程中多步手工处理和反复优化的依赖。",
        "Dream3R 当前使用的教师候选主要以点图和置信度形式进入融合模块。置信度不是最终结论，只是候选质量的一个输入信号。最终权重仍由融合模块结合状态和冲突情况计算。":
            "Dream3R 当前使用的教师候选主要以点图和置信度形式进入融合模块。置信度只表示候选质量的一部分，最终权重仍由融合模块结合状态和冲突情况计算。",
        "Dream状态表示当前重建任务中的状态信号，可以理解为融合时需要参考的当前条件。记忆状态更偏向历史上下文，记录前面窗口、前面视角或已有候选结果带来的信息。二者都不是最终几何，也不替代教师模型。":
            "Dream 状态表示当前重建任务中的状态信号，可以理解为融合时需要参考的当前条件。记忆状态更偏向历史上下文，记录前面窗口、前面视角或已有候选结果带来的信息。二者都作为融合判断的辅助输入使用。",
        "在 Dream3R 中，状态信息进入候选权重计算过程。为了证明它不是形式上的输入，实验设置了正常状态、无状态和乱序状态三组对照。只有正常状态稳定优于后两者，才说明状态信息在当前融合中有用。":
            "在 Dream3R 中，状态信息进入候选权重计算过程。实验设置了正常状态、无状态和乱序状态三组对照，用来判断状态信号是否真正影响融合结果。正常状态稳定优于后两者时，说明状态信息在当前融合中有用。",
        "在 Dream3R 中，它的作用可以概括为筛选有效上下文。它不直接输出最终点图，而是帮助状态分支和候选融合模块减少无关信息影响，尤其适合处理多视角窗口中相邻帧、局部区域和候选片段之间的关系。":
            "在 Dream3R 中，它的作用可以概括为筛选有效上下文。它帮助状态分支和候选融合模块减少无关信息影响，尤其适合处理多视角窗口中相邻帧、局部区域和候选片段之间的关系。",
        "在 Dream3R 中，Mamba 的位置是序列状态建模。它关注输入窗口前后状态如何变化，而不是直接替代教师模型做几何预测。这样可以把历史上下文压缩成对融合有用的状态信号，供后续候选权重计算使用。":
            "在 Dream3R 中，Mamba 的位置是序列状态建模。它关注输入窗口前后状态如何变化，把历史上下文压缩成对融合有用的状态信号，供后续候选权重计算使用。",
        "这一路线保留了教师模型的几何能力，并在最后一步实现了可检查的融合。Dream3R 不负责替代所有教师模型，而是负责在已有候选之间做更稳的组合。":
            "这一路线保留了教师模型的几何能力，并在最后一步实现可检查的融合。Dream3R 的任务是在已有候选之间做更稳的组合。",
        "候选几何分支接收不同教师模型的点图输出。当前主要候选包括 Fast3R、MASt3R、Spann3R 和 VGGT-Omega。候选几何分支的任务是把不同来源的点图和置信度整理成统一张量，供后面的融合模块读取。它不直接决定最终输出，而是提供候选的点图和置信度。":
            "候选几何分支接收不同教师模型的点图输出。当前主要候选包括 Fast3R、MASt3R、Spann3R 和 VGGT-Omega。候选几何分支的任务是把不同来源的点图和置信度整理成统一张量，供后面的融合模块读取，为最终融合提供候选的点图和置信度。",
        "状态条件融合模块是 Dream3R 的核心。它接收候选点图、候选置信度、状态信息和冲突分数，计算每个候选在当前样本中的权重。权重不按固定表格给出，也不只按置信度排序，而是结合候选之间是否一致、状态信息是否支持当前候选、以及不同候选在局部区域的可靠程度来确定。":
            "状态条件融合模块是 Dream3R 的核心。它接收候选点图、候选置信度、状态信息和冲突分数，计算每个候选在当前样本中的权重。权重会综合候选之间是否一致、状态信息是否支持当前候选、以及不同候选在局部区域的可靠程度来确定。",
        "两个数据域中，正常状态均取得最低误差。KITTI 上正常状态比无状态和乱序状态更低；ETH3D 上正常状态也优于两组对照。这说明状态信息不是只作为输入摆在结构里，而是在当前融合设置下影响了候选权重和最终结果。":
            "两个数据域中，正常状态均取得最低误差。KITTI 上正常状态比无状态和乱序状态更低；ETH3D 上正常状态也优于两组对照。这说明状态信息在当前融合设置下影响了候选权重和最终结果。",
    }
    for p in doc.paragraphs:
        text = p.text.strip()
        if text in replacements:
            p.text = replacements[text]


def main():
    doc = Document(SRC)
    polish_theory_sections(doc)
    add_module_sections(doc)
    remove_stiff_contrast_phrasing(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
